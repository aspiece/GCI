from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------------------------------------
# CONFIGURATION
# ----------------------------------------------------------
LABS = [
    (1, "Tool Familiarization & ESD Safety", "Identify and safely use technician tools, ESD gear, and safety procedures."),
    (2, "Internal PC Components & Ports Identification", "Identify motherboard sockets, RAM, storage interfaces, and power connectors."),
    (3, "Assembling a Desktop Computer", "Full PC build from parts; verify POST and BIOS/UEFI detection."),
    (4, "BIOS/UEFI Configuration & Firmware Updates", "Navigate UEFI, set boot order, enable virtualization, update firmware."),
    (5, "Storage Device Installation & Configuration", "Install SATA and NVMe drives; partition, format, and verify functionality."),
    (6, "Memory & CPU Installation and Testing", "Seat and test RAM modules, install CPU and cooling solution."),
    (7, "Power Supply Testing & Cable Management", "Test voltages with a multimeter; verify PSU wattage and proper routing."),
    (8, "Peripheral & Printer Setup", "Install and configure wired/wireless printers, input, and output devices."),
    (9, "Networking & Cable Creation", "Crimp RJ-45 cables, build a small LAN, configure IP settings, and test."),
    (10, "Mobile Device Disassembly & Maintenance", "Open a laptop or tablet safely; identify internal components and perform cleaning or RAM/SSD replacement."),
    (11, "Operating System Installation (Windows and Linux)", "Install, update, and verify dual-boot or virtual OS environments."),
    (12, "System Configuration & User Account Management", "Manage accounts, permissions, startup services, and recovery tools."),
    (13, "System Maintenance & Performance Monitoring", "Use Task Manager, Event Viewer, and maintenance utilities to monitor and tune systems."),
    (14, "Security Hardening & Malware Removal", "Configure antivirus, firewall, encryption, and practice safe malware removal steps."),
    (15, "Capstone – PC Build & Troubleshooting Challenge", "Simulate hardware and OS faults; students diagnose, repair, and document their process.")
]

NUM_LABS = len(LABS)
OUTPUT_FILE = "GCI_Computer_Hardware_Labs_Teacher_Edition_updated.docx"

# Per-lab specific materials and task lists. Each entry is a list of (item, qty, total)
LAB_DETAILS = {
    1: {
        "materials": [
            ("ESD wrist strap", "1 per student", "15"),
            ("Technician tool kit (screwdrivers, pliers)", "1 per student", "15"),
            ("Safety goggles", "1 per student", "15")
        ],
        "tasks": [
            "Review ESD theory and hazards.",
            "Demonstrate proper grounding and strap usage.",
            "Identify common technician tools and safe usage.",
            "Complete the ESD safety checklist worksheet."
        ]
    },
    2: {
        "materials": [
            ("Motherboard poster / diagram", "1", "1"),
            ("Loose sample components (RAM, drives, cards)", "Sets", "5"),
            ("Labeling stickers", "bulk", "1")
        ],
        "tasks": [
            "Label motherboard sockets, ports, and headers.",
            "Match sample components to their connectors.",
            "Complete a ports identification quiz.",
        ]
    },
    3: {
        "materials": [
            ("Computer cases", "1 per group", "5"),
            ("Power supply units (PSU)", "1 per group", "5"),
            ("Motherboards, CPUs, RAM, storage, GPU (as available)", "Sets", "5"),
            ("Thermal paste", "tubes", "3")
        ],
        "tasks": [
            "Assemble a complete desktop from parts.",
            "Verify POST and enter BIOS/UEFI to check device detection.",
            "Document cable routing and basic cable management."
        ]
    },
    4: {
        "materials": [
            ("USB flash drives with firmware utilities", "1 per group", "5"),
            ("Manufacturer BIOS/UEFI update notes", "1", "1"),
            ("Firmware release changelog (print)", "1", "1")
        ],
        "tasks": [
            "Backup current BIOS/UEFI settings.",
            "Perform a firmware update in a controlled environment.",
            "Document the process and any changes made to settings."
        ]
    },
    5: {
        "materials": [
            ("SATA HDD/SSD", "per group", "5"),
            ("NVMe SSDs / M.2 adapters", "sets", "5"),
            ("SATA cables", "various", "10")
        ],
        "tasks": [
            "Install SATA and NVMe devices properly.",
            "Partition and format drives (Windows & Linux examples).",
            "Run a basic read/write verification."
        ]
    },
    6: {
        "materials": [
            ("Spare RAM modules (various sizes)", "sets", "5"),
            ("CPU sockets sample / spare CPU (if available)", "sets", "5"),
            ("Thermal paste and cooler hardware", "sets", "5")
        ],
        "tasks": [
            "Install and seat RAM modules correctly.",
            "Install CPU and cooler, apply thermal paste.",
            "Run a simple memory/CPU stability check (memtest/CPU burn)."
        ]
    },
    7: {
        "materials": [
            ("Multimeter", "1 per group", "5"),
            ("PSU tester (or spare PSU)", "1 per group", "5"),
            ("Cable ties and routing supplies", "bulk", "1")
        ],
        "tasks": [
            "Measure PSU voltages and verify rails within tolerance.",
            "Inspect and document cable routing and airflow considerations.",
            "Practice secure cable fastening and labeling."
        ]
    },
    8: {
        "materials": [
            ("Networked printer or USB printer", "1 per group", "5"),
            ("Printer drivers (USB/Network)", "USB/links", "1"),
            ("Standard peripherals (keyboard/mouse)", "sets", "5")
        ],
        "tasks": [
            "Install and configure a printer (local and network).",
            "Demonstrate driver installation and basic troubleshooting.",
            "Document successful print jobs and common errors."
        ]
    },
    9: {
        "materials": [
            ("RJ-45 crimper", "1 per group", "5"),
            ("Cat5e/Cat6 bulk cable", "rolls", "2"),
            ("Cable tester", "1 per group", "5")
        ],
        "tasks": [
            "Cut, strip, and crimp RJ-45 connectors to make patch cables.",
            "Test cables for continuity and correct wiring.",
            "Configure IP addresses and verify network connectivity."
        ]
    },
    10: {
        "materials": [
            ("Laptop toolkit", "1 per group", "5"),
            ("Replacement SSD/RAM (common laptop sizes)", "sets", "5"),
            ("Compressed air and cleaning brushes", "sets", "5")
        ],
        "tasks": [
            "Safely disassemble a laptop to access internals.",
            "Perform cleaning and replace RAM or SSD where applicable.",
            "Reassemble and verify proper boot and device recognition."
        ]
    },
    11: {
        "materials": [
            ("Windows and Linux install media (USB)", "1 per group", "5"),
            ("Drivers / driver packs", "links/USB", "1"),
            ("Virtualization software (optional)", "1 per lab", "1")
        ],
        "tasks": [
            "Install an OS (Windows) from USB and complete initial setup.",
            "Install a Linux distribution (or VM) and verify hardware support.",
            "Document partitioning choices and bootloader configuration."
        ]
    },
    12: {
        "materials": [
            ("Admin account checklist", "print", "1"),
            ("Sample user accounts to create", "list", "1"),
            ("Recovery media (USB)", "1 per group", "5")
        ],
        "tasks": [
            "Create and manage local user accounts and groups.",
            "Configure permissions and UAC settings.",
            "Set up recovery options and document rollback steps."
        ]
    },
    13: {
        "materials": [
            ("Performance monitoring guide", "print", "1"),
            ("System utilities (Task Manager, PerfMon)", "installed", "1"),
            ("Sample baseline test scripts", "print/USB", "1")
        ],
        "tasks": [
            "Capture baseline performance metrics.",
            "Use Event Viewer to find and interpret system logs.",
            "Make recommendations to improve performance based on findings."
        ]
    },
    14: {
        "materials": [
            ("Antivirus tools (trial or lab versions)", "links", "1"),
            ("Isolated malware removal VM (safe samples)", "1 lab VM", "1"),
            ("Backup and restore media", "USB", "1")
        ],
        "tasks": [
            "Scan and remove simulated malware in an isolated environment.",
            "Configure firewall/antivirus baseline policies.",
            "Document the remediation steps taken."
        ]
    },
    15: {
        "materials": [
            ("Full component kits for builds", "sets", "5"),
            ("Troubleshooting scenario cards", "deck", "1"),
            ("Assessment rubric", "print", "1")
        ],
        "tasks": [
            "Work through assigned fault scenarios and diagnose the cause.",
            "Repair or reconfigure hardware/software to return system to working order.",
            "Prepare a short report documenting steps, root cause, and lessons learned."
        ]
    }
}

# Per-lab exemplar responses for reflections: (quick, applied, deep)
EXAMPLARS = {
    1: (
        "Exemplar: Always wear an ESD strap and connect it to a grounded point before handling components.",
        "Exemplar: A tech who skipped grounding installed RAM that later failed when the system wouldn't POST, causing lost class time and replacement costs.",
        "Exemplar: Following ESD protocols demonstrates professional care, reduces replacement costs, and builds trust with clients."
    ),
    2: (
        "Exemplar: Identify ports by name and purpose (e.g., USB-A for peripherals, HDMI for displays).",
        "Exemplar: Misidentifying a power connector and forcing it into the wrong header can short a board and require replacement.",
        "Exemplar: Accurate identification reduces repair time and mistakes, reinforcing professional workmanship."
    ),
    3: (
        "Exemplar: Confirm POST and BIOS/UEFI device detection after assembling components.",
        "Exemplar: Skipping POST checks can hide an unseated CPU cooler, leading to overheating during testing.",
        "Exemplar: Systematic assembly and verification improves reliability and student troubleshooting skills."
    ),
    4: (
        "Exemplar: Record current firmware settings and create a recovery USB before updating.",
        "Exemplar: An interrupted firmware update can brick a board; a prepared recovery plan prevented data loss in a past incident.",
        "Exemplar: Controlled firmware updates maintain system stability and teach risk mitigation."
    ),
    5: (
        "Exemplar: Verify drive recognition in BIOS/OS after installation and before partitioning.",
        "Exemplar: Installing the wrong drive into a RAID array without verifying can result in degraded array performance or data loss.",
        "Exemplar: Careful disk handling and verification protects student work and demonstrates best practices."
    ),
    6: (
        "Exemplar: Seat RAM until the latches click and apply even pressure when installing a CPU cooler.",
        "Exemplar: Improper cooler mounting caused thermal throttling in a demo system until the mount was corrected.",
        "Exemplar: Proper hardware installation reduces failures and fosters confidence in maintenance tasks."
    ),
    7: (
        "Exemplar: Use a multimeter to check PSU voltages on all rails before connecting critical components.",
        "Exemplar: Identifying a failing 12V rail early prevented damage to GPUs in an earlier lab setup.",
        "Exemplar: Accurate testing improves diagnostic skills and prevents downstream hardware faults."
    ),
    8: (
        "Exemplar: Install printer drivers and verify a test page prints from a networked PC.",
        "Exemplar: A network misconfiguration prevented printing until the DNS and driver settings were corrected.",
        "Exemplar: Peripheral setup is a common support task that builds real-world troubleshooting experience."
    ),
    9: (
        "Exemplar: Confirm cable wiring matches T568B and test continuity with a cable tester.",
        "Exemplar: A miswired cable caused intermittent connectivity; replacing and testing fixed the issue quickly.",
        "Exemplar: Proper cable construction ensures reliable network performance and professional installations."
    ),
    10: (
        "Exemplar: Follow manufacturer guides when disassembling a laptop and keep track of screws and connectors.",
        "Exemplar: Cleaning contacts and replacing a failing SSD resolved slow boot times in a student device.",
        "Exemplar: Safe disassembly practices minimize accidental damage and extend device life."
    ),
    11: (
        "Exemplar: Choose appropriate partitions and document the bootloader choices for dual-boot setups.",
        "Exemplar: A missing driver prevented network access until the correct NIC driver was installed.",
        "Exemplar: Proper OS installation and documentation prepare students for deployment tasks."
    ),
    12: (
        "Exemplar: Create least-privilege accounts and verify UAC and recovery options are configured.",
        "Exemplar: An admin account left unprotected allowed unauthorized changes in a test VM; enforcing policies prevented recurrence.",
        "Exemplar: Good account management practices support security and maintain operational integrity."
    ),
    13: (
        "Exemplar: Collect baseline CPU, memory, and disk metrics before making tuning changes.",
        "Exemplar: Event Viewer revealed a recurring driver error that, once fixed, improved system stability.",
        "Exemplar: Regular maintenance and monitoring reduce downtime and teach preventative care."
    ),
    14: (
        "Exemplar: Run scans in an isolated environment and document remediation steps thoroughly.",
        "Exemplar: An isolated malware sample was analyzed and removed without affecting production systems.",
        "Exemplar: Security workflows and containment protect data and teach responsible remediation."
    ),
    15: (
        "Exemplar: Use a structured troubleshooting approach to isolate hardware vs software faults.",
        "Exemplar: Recreating a fault scenario and applying stepwise fixes demonstrated clear root-cause resolution.",
        "Exemplar: Documenting the process strengthens problem-solving skills and provides evidence of learning."
    )
}

# Per-lab extended descriptions and learning objectives
LAB_DESCRIPTIONS = {
    1: "Introduce students to common technician tools, safety procedures, and electrostatic discharge (ESD) prevention techniques.",
    2: "Hands-on identification of internal PC components, ports, and connectors to build hardware literacy.",
    3: "Guide students through assembling a desktop computer from parts, validating system startup, and basic BIOS checks.",
    4: "Teach BIOS/UEFI navigation, configuration changes for common settings, and safe firmware update procedures.",
    5: "Install and configure both SATA and NVMe storage devices and verify their operation in an OS environment.",
    6: "Practice safe installation and testing of memory modules and CPUs, including cooling and stability checks.",
    7: "Demonstrate PSU testing, safe cable management, and methods to verify power delivery to components.",
    8: "Install and configure peripherals and printers on both local and networked systems, including driver management.",
    9: "Create and test Ethernet patch cables, set up a small LAN, and practice basic network configuration and troubleshooting.",
    10: "Safely disassemble and reassemble mobile devices (laptops/tablets), perform maintenance tasks, and replace common internal components.",
    11: "Install operating systems (Windows and Linux), manage partitions, drivers, and bootloader configuration for single or dual-boot setups.",
    12: "Configure system accounts, permissions, and recovery options to practice secure system administration tasks.",
    13: "Use system monitoring tools to capture performance baselines, analyze logs, and recommend performance improvements.",
    14: "Practice malware detection and removal in an isolated environment while implementing baseline security hardening measures.",
    15: "Capstone project: combine hardware assembly, troubleshooting, OS setup, and documentation into a real-world repair/diagnostic challenge."
}

LAB_OBJECTIVES = {
    1: ["Identify and safely use technician tools", "Demonstrate proper ESD grounding", "Complete an ESD safety checklist"],
    2: ["Name major motherboard components", "Identify common ports and connectors", "Match components to their interfaces"],
    3: ["Assemble a desktop system", "Verify POST and BIOS device detection", "Document assembly and cable routing"],
    4: ["Navigate BIOS/UEFI settings", "Create a firmware backup and recovery plan", "Safely apply firmware updates"],
    5: ["Install SATA and NVMe drives", "Partition and format drives", "Verify read/write operation"],
    6: ["Install RAM and CPU correctly", "Apply thermal paste and mount cooling", "Run basic stability tests"],
    7: ["Measure PSU outputs with a multimeter", "Verify proper cable routing and labeling", "Assess PSU capacity vs. system needs"],
    8: ["Install peripherals and drivers", "Configure networked printers", "Troubleshoot common peripheral issues"],
    9: ["Terminate and crimp RJ-45 connectors", "Test cable continuity and wiring", "Configure basic IP settings and test connectivity"],
    10: ["Safely disassemble a laptop", "Replace or upgrade RAM/SSD", "Reassemble and verify proper boot"],
    11: ["Install Windows/Linux from USB", "Install drivers and update the OS", "Document partitioning and bootloader choices"],
    12: ["Create and manage user accounts", "Configure permissions and UAC", "Set up recovery media and options"],
    13: ["Capture performance metrics", "Analyze Event Viewer logs", "Recommend and apply tuning steps"],
    14: ["Scan and contain malware in an isolated VM", "Apply baseline security settings", "Document remediation steps"],
    15: ["Diagnose hardware/software faults", "Apply repairs and verify operation", "Produce a final report documenting root cause and fixes"]
}

# Default materials/tasks used when a lab has no specific entry

# Per-lab assessment rubrics (criterion, points)
RUBRICS = {
    1: [("Safety & ESD compliance", 8), ("Tool handling & identification", 7), ("Worksheet accuracy", 5), ("Professional conduct", 5)],
    2: [("Identification accuracy", 10), ("Matching components", 6), ("Quiz score", 5), ("Participation", 4)],
    3: [("Assembly completeness", 10), ("POST/BIOS verification", 7), ("Cable management", 4), ("Documentation", 4)],
    4: [("Backup & rollback plan", 6), ("Firmware update procedure", 10), ("Change documentation", 5), ("Safety precautions", 4)],
    5: [("Correct installation", 8), ("Partition/format accuracy", 7), ("Verification testing", 6), ("Cleanup & labeling", 4)],
    6: [("RAM/CPU seating technique", 8), ("Thermal application & mounting", 7), ("Stability testing", 6), ("Documentation", 4)],
    7: [("Voltage measurements accuracy", 10), ("PSU troubleshooting", 6), ("Cable routing & safety", 5), ("Reporting", 4)],
    8: [("Printer installation", 8), ("Driver/configuration", 6), ("Troubleshooting steps", 6), ("User documentation", 5)],
    9: [("Cable construction quality", 10), ("Wiring correctness", 6), ("Connectivity test", 5), ("Network config", 4)],
    10: [("Disassembly/Reassembly care", 8), ("Component replacement", 7), ("Boot verification", 6), ("Cleanliness & documentation", 4)],
    11: [("OS install correctness", 10), ("Driver/install troubleshooting", 6), ("Partition/bootloader setup", 5), ("Documentation", 4)],
    12: [("Account & permissions setup", 8), ("Recovery options configured", 6), ("Policy application", 6), ("Documentation", 5)],
    13: [("Baseline capture accuracy", 8), ("Log analysis", 7), ("Tuning recommendations", 6), ("Report quality", 4)],
    14: [("Malware detection/removal", 10), ("Policy/hardening steps", 7), ("Isolation & safety", 4), ("Remediation report", 4)],
    15: [("Diagnosis accuracy", 10), ("Repair effectiveness", 8), ("Time management/problem solving", 4), ("Final report & lessons learned", 3)]
}

# ----------------------------------------------------------
# HELPER FUNCTIONS
# ----------------------------------------------------------
def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Genesee Career Institute\nComputer Hardware & CompTIA A+ Certification Labs"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.bold = True
        run.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.text = "GCI Computer Hardware Labs – Teacher Edition | Page "
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_materials_table(doc, materials):
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Qty per Group"
    hdr[2].text = "Total (Est.)"
    for item, qty, total in materials:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = item, qty, total
    doc.add_paragraph(" ")

def add_tasks_and_reflections(doc, tasks):
    doc.add_heading("Procedure and Reflection", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Tasks / Procedure", "Reflections"
    row = table.add_row().cells
    row[0].text = "\n".join(tasks)
    right = row[1]
    # Reflection prompts + exemplar responses (exemplar shown in teal and italic)
    # The function accepts lab-specific exemplar text via a global EXAMPLARS mapping keyed by lab number.
    # If no lab_num provided or no exemplar found, fall back to a generic exemplar set.
    lab_num = None
    # detect if caller attached lab number via an attribute on the doc (convention)
    if hasattr(doc, "_current_lab_num"):
        lab_num = doc._current_lab_num

    generic = (
        "Exemplar: To prevent electrostatic discharge that can damage sensitive components; the strap grounds the technician and protects hardware.",
        "Exemplar: A technician replaced a RAM module without grounding and the module failed POST; the faulty replacement required a costly part swap and downtime.",
        "Exemplar: Strict ESD procedures show professionalism, reduce customer costs by preventing avoidable failures, and help build trust in service quality."
    )
    exs = EXAMPLARS.get(lab_num, generic)

    # Insert prompts and place each exemplar directly below its corresponding prompt
    prompts = ["Quick Check (Short Answer):", "Applied Scenario (Medium Response):", "Deep Reflection (Extended):"]
    for idx, prompt in enumerate(prompts):
        p = right.add_paragraph(prompt)
        # pull the matching exemplar (if available)
        exemplar_text = None
        if idx < len(exs):
            exemplar_text = exs[idx]
        if exemplar_text:
            exp = right.add_paragraph(exemplar_text)
            for run in exp.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 102, 153)
    doc.add_paragraph(" ")

def add_teacher_notes(doc, notes, obj_code):
    doc.add_heading("Teacher Notes & Grading Checklist", level=2)
    for note in notes:
        doc.add_paragraph("✔ " + note)
    para = doc.add_paragraph(f"Aligned Objective: {obj_code}")
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Note: do not add a page break here so the assessment rubric remains with the lab

# ----------------------------------------------------------
# MAIN DOCUMENT
# ----------------------------------------------------------
doc = Document()
# Use narrow (half-inch) margins for a compact teacher edition
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
add_header_footer(doc)

# Title page content
doc.add_heading("CompTIA A+ Computer Hardware Labs", level=1)
doc.add_paragraph("Teacher Edition\nGenesee Career Institute\n\n")

# Table of Contents (insert a live TOC field - Word will populate it when fields are updated)
p = doc.add_paragraph()
try:
    # Add a Word TOC field. This uses python-docx's underlying XML helper; Word will populate it when fields are updated (select field and press F9).
    p._p.add_fldSimple('TOC \\o "1-3" \\h \\z \\u')
except Exception:
    # Fallback: visible placeholder if fldSimple isn't available in this environment
    p.add_run("Table of Contents (Update in Word to activate hyperlinks)")

# Ensure the TOC field is on its own page
doc.add_page_break()

# Insert the lab list with brief summaries on the following page for quick reference.
doc.add_heading("Lab Titles & Summaries", level=2)
for num, title, summary in LABS:
    # Use a bold inline title followed by a short summary to appear on the TOC-reference page
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {title}")
    run.bold = True
    p.add_run(f" — {summary}")
doc.add_page_break()

# NOTE: python-docx cannot reliably insert a live Word Table of Contents field.
# To create a working TOC after opening this document in Microsoft Word:
# 1) Place the cursor on the 'Table of Contents' page, or select the TOC area.
# 2) Go to References → Table of Contents → Choose a style, or press Ctrl+A then F9 to update fields.
# 3) Save the document. This will populate an automatic TOC based on the heading styles used in this document.

# Lab Overview
doc.add_heading("Lab Sequence Overview", level=1)
doc.add_paragraph("Core 1 – Hardware & Networking: Labs 1–10\nCore 2 – Operating Systems & Security: Labs 11–15")
doc.add_paragraph("Below is a concise overview of each lab with its focus and expected student outcomes.")
for num, title, summary in LABS:
    doc.add_heading(f"Lab {num} — {title}", level=2)
    doc.add_paragraph(summary)
doc.add_page_break()

# Materials (simple sample)
doc.add_heading("Master Materials List (By Type)", level=1)
materials = [
    ("ESD mat and wrist strap", "1", "5"),
    ("Screwdriver set", "1", "5"),
    ("Cable tester", "1", "2"),
    ("Multimeter", "1", "2"),
    ("Sample PC components", "Set", "5")
]
add_materials_table(doc, materials)
doc.add_page_break()

# Generate all labs with per-lab materials and tasks where available
default_materials = [
    ("ESD mat and wrist strap", "1", "5"),
    ("Screwdriver set", "1", "5"),
    ("Cable tester", "1", "2"),
    ("Multimeter", "1", "2"),
    ("Sample PC components", "Set", "5")
]

default_tasks = [
    "1. Follow lab setup instructions provided by the instructor.",
    "2. Perform the required configuration or assembly task.",
    "3. Verify proper function and document results.",
    "4. Capture photo evidence of your completed work."
]

for i in range(1, NUM_LABS + 1):
    num, title, summary = LABS[i - 1]
    # ensure each lab starts on its own page; first lab follows the previous page break
    if i > 1:
        doc.add_page_break()
    doc.add_heading(f"Lab {num} – {title}", level=1)
    doc.add_paragraph(summary)
    doc.add_paragraph("⏱ Estimated Time: 90 minutes")
    doc.add_paragraph("Core 1 – Hardware & Networking" if i <= 10 else "Core 2 – Operating Systems & Security")
    # Insert a short description and a bulleted list of learning objectives (per-lab)
    doc.add_heading("Description", level=2)
    desc = LAB_DESCRIPTIONS.get(i, summary)
    doc.add_paragraph(desc)

    doc.add_heading("Learning Objectives", level=2)
    objectives = LAB_OBJECTIVES.get(i, ["Complete the hands-on tasks and apply best practices for computer hardware assembly and troubleshooting."])
    for obj in objectives:
        # use the 'List Bullet' style for a simple bulleted list
        try:
            doc.add_paragraph(obj, style='List Bullet')
        except Exception:
            # fallback to plain paragraph with a bullet prefix
            doc.add_paragraph(f"• {obj}")

    # Materials: prefer LAB_DETAILS entry, otherwise fall back to defaults
    lab_materials = LAB_DETAILS.get(i, {}).get("materials", default_materials)
    doc.add_heading("Materials Needed", level=2)
    add_materials_table(doc, lab_materials)

    # Tasks and reflections: prefer LAB_DETAILS tasks if present
    lab_tasks = LAB_DETAILS.get(i, {}).get("tasks", default_tasks)
    # attach current lab number to doc so add_tasks_and_reflections can pick lab-specific exemplars
    doc._current_lab_num = i
    add_tasks_and_reflections(doc, lab_tasks)
    # remove the temporary attribute to avoid leakage
    delattr(doc, "_current_lab_num")

    add_teacher_notes(doc,
        ["Ensure each student has safety gear and workstation ready.",
         "Observe procedure compliance and tool handling.",
         "Award points for accuracy, documentation, and teamwork."],
        f"CompTIA A+ 220-120{1 if i <= 10 else 2} | 5.{i if i <= 10 else i - 10}"
    )
    # Assessment rubric (show points and percent of lab total)
    rubric = RUBRICS.get(i, [("Completion", 10), ("Accuracy", 10), ("Documentation", 5)])
    doc.add_heading("Assessment Rubric", level=2)
    rtable = doc.add_table(rows=1, cols=3)
    rh = rtable.rows[0].cells
    rh[0].text = "Criterion"
    rh[1].text = "Points"
    rh[2].text = "Percent"
    total_points = 0
    for crit, pts in rubric:
        total_points += pts
    # add rows with percent column
    for crit, pts in rubric:
        row = rtable.add_row().cells
        row[0].text = crit
        row[1].text = str(pts)
        percent = (pts / total_points * 100) if total_points else 0
        # round to nearest whole percent for clarity
        row[2].text = f"{int(round(percent))}%"
    doc.add_paragraph(f"Total points: {total_points}")

# Save document
# Add a condensed rubric summary (one section with per-lab totals and criterion percentages)
doc.add_page_break()
doc.add_heading("Rubric Summary (Condensed)", level=1)
# Create a compact multi-column table to force a condensed, printable layout.
num_cols = 3
cols = num_cols
per_col = (len(LABS) + cols - 1) // cols
rtable = doc.add_table(rows=1, cols=cols)
rtable.autofit = False
# Optionally set column widths to distribute across the page (approximate)
try:
    widths = [Inches(6.5 / cols)] * cols
    for ci, w in enumerate(widths):
        rtable.columns[ci].width = w
except Exception:
    pass

for col in range(cols):
    cell = rtable.rows[0].cells[col]
    # Build lab entries assigned to this column
    start = col * per_col
    end = min(start + per_col, len(LABS))
    for idx in range(start, end):
        num, title, _ = LABS[idx]
        rubric = RUBRICS.get(num, [])
        total = sum(p for _, p in rubric) if rubric else 0
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        title_run = p.add_run(f"Lab {num} — {title}")
        title_run.bold = True
        title_run.font.size = Pt(9)
        total_run = p.add_run(f" — Total: {total} pts")
        total_run.font.size = Pt(9)
        if rubric:
            for crit, pts in rubric:
                pct = int(round((pts / total * 100) if total else 0))
                r = cell.add_paragraph(f"• {crit} — {pts} pts ({pct}%)")
                r.paragraph_format.space_before = Pt(0)
                r.paragraph_format.space_after = Pt(0)
                for run in r.runs:
                    run.font.size = Pt(8.5)

doc.save(OUTPUT_FILE)
print(f"✅ Lab Manual generated: {OUTPUT_FILE}")
