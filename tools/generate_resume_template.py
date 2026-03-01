from pathlib import Path

resume_out_path = Path(
    r"c:\github\GCI\wordprojectfiles\Lastname_Firstname_Resume_Template.docx"
)
cover_letter_out_path = Path(
    r"c:\github\GCI\wordprojectfiles\Lastname_Firstname_CoverLetter_Template.docx"
)

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING


def add_heading_and_text(document, heading, text):
    document.add_heading(heading, level=1)
    document.add_paragraph(text)


def apply_accessible_template_formatting(document):
    for section in document.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    normal_paragraph = normal_style.paragraph_format
    normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_paragraph.line_spacing = 1.15
    normal_paragraph.space_after = Pt(6)

    heading1_style = document.styles["Heading 1"]
    heading1_style.font.name = "Calibri"
    heading1_style.font.size = Pt(14)
    heading1_style.paragraph_format.space_before = Pt(10)
    heading1_style.paragraph_format.space_after = Pt(4)

    title_style = document.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(18)


resume = Document()
apply_accessible_template_formatting(resume)
resume.add_heading("Career Essentials Resume Template", level=0)
resume.add_paragraph("Replace this sample text with your own real information.")

resume.add_paragraph(
    "Maya Johnson | maya.johnson@email.com | (555) 018-2034 | Charlotte, NC"
)
resume.add_paragraph("Target Role: Customer Support Intern")

add_heading_and_text(
    resume,
    "Professional Summary",
    "Service-focused high school senior seeking a Customer Support Intern role. Strong in customer communication, ticket documentation, and issue resolution with experience supporting students in a fast-paced school help desk environment.",
)

add_heading_and_text(
    resume,
    "Core Skills",
    "Customer Communication; Technical Troubleshooting; Ticket Documentation; Google Workspace; Data Entry; Time Management; Team Collaboration; Problem Solving",
)

add_heading_and_text(
    resume,
    "Education",
    "Westbrook High School, Charlotte, NC — Expected Graduation 2027\nRelevant Coursework: Business Technology, Computer Applications, Professional Communication",
)

resume.add_heading("Experience", level=1)
resume.add_paragraph(
    "Student Tech Help Desk Volunteer | Westbrook High School | 2025–Present"
)
resume.add_paragraph(
    "• Assisted 40+ students and staff with login, Chromebook, and account-access issues, improving same-day resolution times during peak periods."
)
resume.add_paragraph(
    "• Logged support requests and fixes in a shared tracker to keep records accurate and improve handoff between team members."
)
resume.add_paragraph(
    "• Explained technical steps in clear, user-friendly language and maintained a professional, calm tone during high-volume support periods."
)

resume.add_paragraph(
    "School Event Team Member | Westbrook HS Activities Office | 2024–2025"
)
resume.add_paragraph(
    "• Coordinated check-in for a 120-student event and reduced wait times by organizing alphabetical lines and digital attendance verification."
)
resume.add_paragraph(
    "• Collaborated with 6 team members to manage setup, attendee support, and post-event follow-up communications."
)

resume.add_paragraph(
    "Community Volunteer Tutor | Charlotte Youth Center | 2024–Present"
)
resume.add_paragraph(
    "• Supported middle-school students with weekly homework and basic computer-skills practice in one-on-one and small-group sessions."
)
resume.add_paragraph(
    "• Tracked attendance and progress notes to provide clear updates to program staff."
)

resume.save(resume_out_path)
print(f"Updated: {resume_out_path}")

cover = Document()
apply_accessible_template_formatting(cover)
cover.add_heading("Career Essentials Cover Letter Template", level=0)
cover.add_paragraph("Replace this sample text with your own real information.")
cover.add_paragraph("Date: ________________________________")
cover.add_paragraph("Hiring Manager")
cover.add_paragraph("Company Name")
cover.add_paragraph("Company Address")
cover.add_paragraph("")

cover.add_paragraph("Dear Hiring Manager,")
cover.add_paragraph(
    "I am excited to apply for the Customer Support Intern position at BrightPath Solutions. In my student tech help desk role, I support over 40 students and staff with account and device issues while maintaining clear communication and accurate documentation. I am especially interested in this role because your posting emphasizes customer communication, troubleshooting, and teamwork—areas I use every week."
)
cover.add_paragraph(
    "Your posting highlights the need for someone who can work effectively in a fast-paced team environment and communicate clearly with users. Through event support and peer tech assistance, I have built those exact skills while staying calm under pressure and solving problems quickly. For example, during our school technology rollout, I helped triage login problems, documented recurring issues, and shared patterns with staff so we could prevent repeat errors."
)
cover.add_paragraph(
    "In addition to technical support, I bring strong organization and follow-through. As a school event team member, I helped coordinate check-in for a 120-student event, which required accurate tracking, clear directions for participants, and consistent communication across our team. I would bring the same reliability, customer-first mindset, and attention to detail to BrightPath Solutions."
)
cover.add_paragraph(
    "I would welcome the opportunity to bring my organization, communication, and service mindset to your team. Thank you for your time and consideration. I look forward to discussing next steps."
)
cover.add_paragraph("")
cover.add_paragraph("Sincerely,")
cover.add_paragraph("Maya Johnson")

cover.save(cover_letter_out_path)
print(f"Updated: {cover_letter_out_path}")
