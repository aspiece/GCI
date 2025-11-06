from docx import Document
 
# Create a new Word document
doc = Document()
doc.add_heading('🧹 Helpdesk Daily Checkout List', level=1)
 
doc.add_paragraph("""
All stations should be tidied, sanitized, and organized before end of day.
Check off each area once it meets the following standards:
• Surfaces cleared of clutter
• Tools/equipment put away
• Trash removed
• Wipes used on high-touch areas
• Power off (if applicable)
""")
 
# Coding Stations
doc.add_heading('🖥️ Coding Stations', level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Station'
hdr_cells[1].text = 'Cleared'
hdr_cells[2].text = 'Wiped'
hdr_cells[3].text = 'Power Off'
hdr_cells[4].text = 'Notes'
 
for i in range(1, 4):
    row_cells = table.add_row().cells
    row_cells[0].text = f'Coding Station {i}'
    row_cells[1].text = '☐'
    row_cells[2].text = '☐'
    row_cells[3].text = '☐'
    row_cells[4].text = ''
 
# 3D Printing Area
doc.add_heading('🖨️ 3D Printing Area', level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Printer'
hdr_cells[1].text = 'Bed Cleared'
hdr_cells[2].text = 'Tools Put Away'
hdr_cells[3].text = 'Power Off'
hdr_cells[4].text = 'Notes'
 
for i in range(1, 6):
    row_cells = table.add_row().cells
    row_cells[0].text = f'R2-D{i}'
    row_cells[1].text = '☐'
    row_cells[2].text = '☐'
    row_cells[3].text = '☐'
    row_cells[4].text = ''
 
row_cells = table.add_row().cells
row_cells[0].text = 'Printing Station'
row_cells[1].text = '☐'
row_cells[2].text = '☐'
row_cells[3].text = '☐'
row_cells[4].text = ''
 
# Technician Stations
doc.add_heading('🧰 Technician Stations', level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Station'
hdr_cells[1].text = 'Cleared'
hdr_cells[2].text = 'Tools Sorted'
hdr_cells[3].text = 'Power Off'
hdr_cells[4].text = 'Notes'
 
for i in range(1, 5):
    row_cells = table.add_row().cells
    row_cells[0].text = f'Technician Station {i}'
    row_cells[1].text = '☐'
    row_cells[2].text = '☐'
    row_cells[3].text = '☐'
    row_cells[4].text = ''
 
# Shared Areas
doc.add_heading('🪑 Shared & Administrative Areas', level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Area'
hdr_cells[1].text = 'Cleared'
hdr_cells[2].text = 'Sanitized'
hdr_cells[3].text = 'Organized'
hdr_cells[4].text = 'Notes'
 
shared_areas = ['Conference Table', 'Manager Desk', 'Assembly Area', 'Emersion Station']
for area in shared_areas:
    row_cells = table.add_row().cells
    row_cells[0].text = area
    row_cells[1].text = '☐'
    row_cells[2].text = '☐'
    row_cells[3].text = '☐'
    row_cells[4].text = ''
 
# End-of-Day Sign-off
doc.add_heading('🧾 End-of-Day Sign-off', level=2)
doc.add_paragraph("""
Checked by: ____________________________
 
Date: _________________________________
 
Notes/Issues: ___________________________________________
""")
 
# Save document
doc.save("Helpdesk_Daily_Checkout_List.docx")
print("✅ Word document created: Helpdesk_Daily_Checkout_List.docx")